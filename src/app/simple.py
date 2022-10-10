import docx
from docx.table import Table
import json

from app.util import JsonObject, jsonpath


def _load_json(json_filename: str) -> JsonObject:
    with open(json_filename, "r", encoding="utf-8") as file:
        return json.load(file)


def _print_vis_rule(elem_json: JsonObject, doc: docx.Document, rules: dict[str, str]):
    vis_id = elem_json["visibilityRuleId"]
    if vis_id:
        vis_rule = rules.get(vis_id, f"external rule {vis_id[:8]}")
        doc.add_paragraph(f"If {vis_rule}:", style="outline_visrule")


class ControlRun(list):
    """List of controls, outwith a table or repeating group"""

    def add_table_to(self, doc):
        table = doc.add_table(rows=0, cols=2, style="outline_2col_table")
        header_row = table.add_row()
        header_row.cells[0].text = "Questions / text"
        header_row.cells[1].text = "Answer type"
        for control in self:
            row = table.add_row()
            row.cells[0].text = jsonpath(control, "label/value") or "???"
            row.cells[1].text = jsonpath(control, "subtype") or "???"

        doc.add_paragraph()


def _parse_subsections(section_json: JsonObject) -> list[JsonObject | ControlRun]:
    """Resolves section into tables, repeating groups, and plain runs of controls"""

    # gather controls from adjoining rows into PlainRun subsections
    subsections = {}
    for row in section_json["rows"]:
        row_num = row["number"]
        subsections[row_num] = subsections.pop(row_num - 1, ControlRun())
        for column in row["columns"]:
            subsections[row_num].extend(column["controls"])

    # gather groups and tables
    for group in section_json["groups"]:
        row_num = group["number"]
        subsections[row_num] = group

    # sort by row number (sequence in form)
    return [subsection for _, subsection in sorted(subsections.items())]


def main(json_filename: str, template_filename: str, output_filename: str):
    if template_filename == output_filename:
        raise ValueError("Output filename should not overwrite template")

    doc = docx.Document(template_filename)
    form_json = _load_json(json_filename)

    rules: dict[str, str] = {
        rule["definitionId"]: rule["name"] for rule in form_json["rules"]
    }

    doc.add_paragraph(form_json["name"], style="outline_formtitle")

    for page_num, page in enumerate(form_json["pages"], start=1):

        _print_vis_rule(page, doc, rules)

        page_title = page["title"]
        doc.add_paragraph(f"Page {page_num}: {page_title}", style="outline_pagetitle")

        for section_num, section in enumerate(page["sections"], start=1):
            _print_vis_rule(section, doc, rules)
            section_title = section["title"]
            doc.add_paragraph(
                f"Section {page_num}.{section_num}: {section_title}",
                style="outline_sectiontitle",
            )
            for subsection in _parse_subsections(section):
                if isinstance(subsection, ControlRun):
                    subsection.add_table_to(doc)
                else:
                    doc.add_paragraph("Group or table")

    doc.save(output_filename)


if __name__ == "__main__":
    main("tests/Lunch planning_2.json", "template.docx", "output.docx")
