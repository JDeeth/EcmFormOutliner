from dataclasses import dataclass
from typing import Optional
import pytest

import docx
from docx.enum.style import WD_STYLE_TYPE

# pylint: disable=missing-function-docstring


@dataclass
class DummyDisplayControl:
    """Same fields as DisplayControl"""

    label: str
    type: str
    details: str = ""
    visibility_rule: str = ""
    mandatory: Optional[bool] = None


@pytest.fixture(name="simple_run")
def fixture_simple_run():
    params = [
        ("Date of meeting ", "Date", "", "", True),
        ("Venue", "Free text"),
    ]
    return [DummyDisplayControl(*p) for p in params]


def add_run_to(doc, run):
    table = doc.add_table(rows=0, cols=2, style="outline_2col_table")

    table.add_row()
    table.rows[-1].cells[0].text = "Questions / text"
    table.rows[-1].cells[1].text = "Answer type"

    for control in run:
        if control.mandatory is True:
            label = f"{control.label}*"
        else:
            label = control.label
        table.add_row()
        table.rows[-1].cells[0].text = label
        table.rows[-1].cells[1].text = control.type


def should_present_plain_runs_as_2col_table(simple_run):
    doc = docx.Document()
    table_style = WD_STYLE_TYPE.TABLE  # pylint: disable=no-member
    doc.styles.add_style("outline_2col_table", style_type=table_style)

    add_run_to(doc, simple_run)

    table = doc.tables[0]
    cells = []
    for row in table.rows:
        cells.append([c.text for c in row.cells])

    assert cells == [
        ["Questions / text", "Answer type"],
        ["Date of meeting *", "Date"],
        ["Venue", "Free text"],
    ]
